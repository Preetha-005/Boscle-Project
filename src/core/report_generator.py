"""
Report generation module for the Meeting Video Captioning Program
Generates comprehensive reports in PDF, Word, and text formats
"""

import os
import time
import json
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any, Union
from datetime import datetime, timedelta
import base64
from io import BytesIO


class ReportGenerator:
    """Generates detailed reports with screenshots, transcripts, and analysis"""
    
    def __init__(self, config_manager, logger):
        self.config_manager = config_manager
        self.logger = logger
        
        # Configuration
        self.report_format = config_manager.get("report_generation.format", "pdf")
        self.include_screenshots = config_manager.get("report_generation.include_screenshots", True)
        self.include_timestamps = config_manager.get("report_generation.include_timestamps", True)
        self.include_summaries = config_manager.get("report_generation.include_summaries", True)
        self.summary_model = config_manager.get("report_generation.summary_model", "openai")
        
        # API credentials - secure from environment variables
        self.openai_api_key = config_manager.get_secure("report_generation.openai_api_key")
        
    def generate_report(self, video_analysis: Dict[str, Any], transcript: Dict[str, Any], 
                       output_dir: str) -> str:
        """
        Generate comprehensive report
        Returns path to the generated report
        """
        start_time = time.time()
        self.logger.info(f"Starting report generation in {self.report_format} format")
        
        try:
            # Create output directory
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Prepare report data
            self.logger.log_processing_step("Preparing report data")
            report_data = self._prepare_report_data(video_analysis, transcript)
            
            # Generate summaries if enabled
            if self.include_summaries:
                self.logger.log_processing_step("Generating content summaries")
                report_data["summaries"] = self._generate_summaries(video_analysis, transcript)
            
            # Generate report based on format
            if self.report_format.lower() == "pdf":
                report_path = self._generate_pdf_report(report_data, output_path)
            elif self.report_format.lower() == "docx":
                report_path = self._generate_docx_report(report_data, output_path)
            elif self.report_format.lower() == "txt":
                report_path = self._generate_text_report(report_data, output_path)
            else:
                # Default to PDF
                report_path = self._generate_pdf_report(report_data, output_path)
                
            # Verify report was created
            if not Path(report_path).exists():
                raise Exception("Report was not created successfully")
                
            processing_time = time.time() - start_time
            self.logger.info(f"Report generation completed in {processing_time:.2f}s: {report_path}")
            
            return report_path
            
        except Exception as e:
            self.logger.error(f"Report generation failed: {str(e)}", exc_info=True)
            raise Exception(f"Failed to generate report: {str(e)}")
            
    def _prepare_report_data(self, video_analysis: Dict[str, Any], 
                           transcript: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare and structure data for report generation"""
        
        # Extract video information
        video_info = video_analysis.get("video_info", {})
        frames = video_analysis.get("frames", [])
        interactions = video_analysis.get("interactions", [])
        analysis_summary = video_analysis.get("summary", {})
        
        # Extract transcript information
        segments = transcript.get("segments", [])
        full_text = transcript.get("full_text", "")
        transcript_stats = transcript.get("statistics", {})
        transcript_summary = transcript.get("summary", {})
        
        # Prepare structured data
        report_data = {
            "metadata": {
                "generation_time": datetime.now().isoformat(),
                "video_file": video_info.get("filename", "Unknown"),
                "video_path": video_info.get("path", ""),
                "report_format": self.report_format,
                "include_screenshots": self.include_screenshots,
                "processing_duration": video_analysis.get("processing_time", 0)
            },
            "video_overview": {
                "filename": video_info.get("filename", "Unknown"),
                "duration": self._format_duration(video_info.get("duration", 0)),
                "resolution": f"{video_info.get('width', 0)}x{video_info.get('height', 0)}",
                "fps": video_info.get("fps", 0),
                "file_size": self._format_file_size(video_info.get("file_size", 0)),
                "codec": video_info.get("codec", "Unknown")
            },
            "analysis_summary": {
                "total_frames_analyzed": analysis_summary.get("total_frames_analyzed", 0),
                "scene_changes_detected": analysis_summary.get("scene_changes", 0),
                "interactions_detected": analysis_summary.get("interactions_detected", 0),
                "timeline_segments": analysis_summary.get("timeline_segments", []),
                "key_moments": analysis_summary.get("key_moments", [])
            },
            "transcript_summary": {
                "total_segments": transcript_stats.get("total_segments", 0),
                "total_words": transcript_stats.get("total_words", 0),
                "total_duration": self._format_duration(transcript_stats.get("total_duration", 0)),
                "average_confidence": transcript_stats.get("average_confidence", 0),
                "words_per_minute": transcript_stats.get("words_per_minute", 0),
                "top_keywords": transcript_summary.get("top_keywords", [])
            },
            "detailed_timeline": self._create_detailed_timeline(frames, interactions, segments),
            "transcript_segments": segments,
            "interactions": interactions,
            "screenshots": self._prepare_screenshots_data(frames) if self.include_screenshots else [],
            "full_transcript": full_text
        }
        
        return report_data
        
    def _create_detailed_timeline(self, frames: List[Dict[str, Any]], 
                                interactions: List[Dict[str, Any]], 
                                segments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Create a unified timeline combining video analysis (excludes transcript - shown separately)"""
        
        timeline_events = []
        
        # Add scene changes
        for frame in frames:
            if frame.get("is_scene_change", False):
                event = {
                    "timestamp": frame["timestamp"],
                    "timestamp_formatted": frame["timestamp_formatted"],
                    "type": "scene_change",
                    "description": f"Scene change detected (confidence: {frame['scene_confidence']:.2f})",
                    "details": frame.get("analysis", {}),
                    "screenshot": frame.get("file_path") if self.include_screenshots else None
                }
                timeline_events.append(event)
                
        # Add interactions
        for interaction in interactions:
            event = {
                "timestamp": interaction["timestamp"],
                "timestamp_formatted": interaction["timestamp_formatted"],
                "type": interaction["type"],
                "description": interaction["description"],
                "confidence": interaction.get("confidence", 0),
                "screenshot": interaction.get("frame_path") if self.include_screenshots else None
            }
            timeline_events.append(event)
        
        # NOTE: Transcript segments are NOT added here - they are shown separately in Full Transcript section
        # This avoids duplication and keeps the timeline focused on visual events
            
        # Sort by timestamp
        timeline_events.sort(key=lambda x: x["timestamp"])
        
        return timeline_events
        
    def _prepare_screenshots_data(self, frames: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Prepare screenshot data for inclusion in report"""
        screenshots = []
        
        for frame in frames:
            if frame.get("file_path") and Path(frame["file_path"]).exists():
                screenshot = {
                    "timestamp": frame["timestamp"],
                    "timestamp_formatted": frame["timestamp_formatted"],
                    "file_path": frame["file_path"],
                    "description": self._generate_frame_description(frame),
                    "is_scene_change": frame.get("is_scene_change", False),
                    "confidence": frame.get("scene_confidence", 0)
                }
                screenshots.append(screenshot)
                
        return screenshots
        
    def _generate_frame_description(self, frame: Dict[str, Any]) -> str:
        """Generate description for a frame based on analysis"""
        analysis = frame.get("analysis", {})
        timestamp = frame.get("timestamp_formatted", "Unknown")
        
        description = f"Frame at {timestamp}"
        
        if frame.get("is_scene_change", False):
            description += " - Scene Change"
            
        content_type = analysis.get("content_type", "unknown")
        if content_type != "unknown":
            description += f" ({content_type})"
            
        if analysis.get("has_text", False):
            description += " - Contains text"
            
        if analysis.get("significant_change", False):
            description += " - Significant visual change"
            
        return description
        
    def _generate_summaries(self, video_analysis: Dict[str, Any], 
                          transcript: Dict[str, Any]) -> Dict[str, Any]:
        """Generate AI-powered summaries of content"""
        summaries = {}
        
        try:
            # Generate overall meeting summary
            full_text = transcript.get("full_text", "")
            if full_text:
                summaries["overall_summary"] = self._generate_text_summary(
                    full_text, "overall meeting content"
                )
                
            # Generate key points summary
            summaries["key_points"] = self._extract_key_points(full_text)
            
            # Generate action items (if any)
            summaries["action_items"] = self._extract_action_items(full_text)
            
            # Generate timeline summary
            interactions = video_analysis.get("interactions", [])
            if interactions:
                summaries["visual_summary"] = self._generate_visual_summary(interactions)
                
        except Exception as e:
            self.logger.warning(f"Summary generation failed: {e}")
            summaries["error"] = f"Could not generate summaries: {str(e)}"
            
        return summaries
        
    def _generate_text_summary(self, text: str, content_type: str) -> str:
        """Generate summary using AI service"""
        
        if self.summary_model == "openai" and self.openai_api_key and "YOUR_" not in self.openai_api_key:
            return self._generate_openai_summary(text, content_type)
        else:
            return self._generate_local_summary(text, content_type)
            
    def _generate_openai_summary(self, text: str, content_type: str) -> str:
        """Generate summary using OpenAI API"""
        try:
            # to be updated by the user - configure OpenAI client
            from openai import OpenAI
            
            client = OpenAI(api_key=self.openai_api_key)
            
            prompt = f"""Please provide a concise summary of the following {content_type}:

{text[:4000]}  # Limit text length for API

Summary should include:
- Main topics discussed
- Key decisions or conclusions
- Important information mentioned

Keep the summary clear and professional."""

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that creates concise, professional summaries of meeting content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.3
            )
            
            return response.choices[0].message.content.strip()
            
        except ImportError:
            self.logger.warning("OpenAI library not available")
            return self._generate_local_summary(text, content_type)
        except Exception as e:
            self.logger.warning(f"OpenAI summary failed: {e}")
            return self._generate_local_summary(text, content_type)
            
    def _generate_local_summary(self, text: str, content_type: str) -> str:
        """Generate a meaningful summary paragraph from the transcript using extractive summarization"""
        if not text or len(text.strip()) < 50:
            return "No summary available - transcript too short."
        
        import re
        
        word_count = len(text.split())
        
        # If transcript is short (under 200 words), just use it as the summary
        if word_count < 200:
            return text.strip()
        
        # Split into sentences properly
        sentences = re.split(r'(?<=[.!?])\s+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]
        
        if not sentences:
            return text.strip()  # Fallback to full text
        
        total_sentences = len(sentences)
        duration_mins = max(1, int(word_count / 150))  # Approximate speaking rate
        
        # Define importance indicators for sentences
        importance_indicators = {
            'decision': ['decided', 'agreed', 'concluded', 'determined', 'resolved', 'confirmed'],
            'action': ['will', 'need to', 'should', 'must', 'going to', 'plan to', 'action', 'next step'],
            'question': ['?', 'how do', 'what is', 'why', 'when will', 'who will'],
            'emphasis': ['important', 'key', 'main', 'critical', 'essential', 'significant', 'focus'],
            'summary': ['in summary', 'to summarize', 'overall', 'in conclusion', 'the goal', 'objective'],
            'topic_intro': ['today we', 'this meeting', 'we\'re here to', 'the purpose', 'agenda']
        }
        
        def score_sentence(sentence, position, total):
            """Score sentence importance based on content and position"""
            score = 0
            lower = sentence.lower()
            
            # Position scoring - first and last sentences are often important
            if position == 0:
                score += 3  # Opening usually sets context
            elif position == total - 1:
                score += 2  # Closing often summarizes
            elif position < total * 0.2:
                score += 1  # Early sentences often introduce key topics
            
            # Content-based scoring
            for category, keywords in importance_indicators.items():
                for keyword in keywords:
                    if keyword in lower:
                        if category == 'decision':
                            score += 4
                        elif category == 'action':
                            score += 3
                        elif category == 'summary':
                            score += 4
                        elif category == 'topic_intro':
                            score += 3
                        elif category == 'emphasis':
                            score += 2
                        elif category == 'question':
                            score += 1
                        break  # Only count once per category
            
            # Longer sentences (within reason) often contain more information
            word_len = len(sentence.split())
            if 10 <= word_len <= 30:
                score += 1
            
            return score
        
        # Score all sentences
        scored = [(s, score_sentence(s, i, total_sentences), i) for i, s in enumerate(sentences)]
        
        # Get top sentences but maintain original order for coherence
        top_scored = sorted(scored, key=lambda x: x[1], reverse=True)[:5]
        # Sort by original position to maintain flow
        key_sentences = sorted(top_scored, key=lambda x: x[2])
        
        # Extract main topics
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        
        # Common stop words to filter
        stop_words = {'this', 'that', 'with', 'have', 'will', 'from', 'they', 'been', 
                     'were', 'being', 'their', 'would', 'could', 'should', 'about',
                     'which', 'there', 'these', 'those', 'just', 'also', 'than',
                     'then', 'more', 'some', 'into', 'over', 'such', 'only', 'other',
                     'very', 'when', 'what', 'your', 'here', 'know', 'think', 'said',
                     'going', 'really', 'actually', 'basically', 'like', 'okay', 'yeah'}
        
        # Count meaningful words
        word_freq = {}
        for word in words:
            if word not in stop_words and len(word) > 4:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top topics (words mentioned multiple times)
        main_topics = [word.title() for word, count in 
                      sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:6] 
                      if count >= 2]
        
        # Build the summary
        summary_parts = []
        
        # Opening - meeting overview
        summary_parts.append(
            f"This is a {duration_mins}-minute meeting recording with approximately {word_count} words spoken."
        )
        
        # Add main topics if found
        if main_topics:
            if len(main_topics) <= 3:
                topics_str = " and ".join(main_topics)
            else:
                topics_str = ", ".join(main_topics[:-1]) + f", and {main_topics[-1]}"
            summary_parts.append(f"Key topics covered include {topics_str}.")
        
        # Add key discussion points from important sentences
        added_content = set()
        discussion_points = []
        
        for sent, score, pos in key_sentences:
            # Clean up the sentence
            clean_sent = sent.strip()
            if len(clean_sent) > 150:
                clean_sent = clean_sent[:147] + "..."
            
            # Avoid duplicate content
            sent_key = clean_sent[:50].lower()
            if sent_key not in added_content and score > 0:
                added_content.add(sent_key)
                discussion_points.append(clean_sent)
        
        # Add discussion points as a flowing narrative
        if discussion_points:
            if len(discussion_points) == 1:
                summary_parts.append(f"A notable point discussed: \"{discussion_points[0]}\"")
            else:
                summary_parts.append("Key points from the discussion include:")
                for i, point in enumerate(discussion_points[:3], 1):
                    summary_parts.append(f"({i}) \"{point}\"")
        
        # Combine into final summary
        full_summary = " ".join(summary_parts)
        
        return full_summary
        
    def _extract_key_points(self, text: str) -> List[str]:
        """Extract key points from transcript"""
        key_points = []
        
        # Look for common key point indicators
        indicators = [
            'important', 'key', 'main', 'primary', 'crucial', 'significant',
            'decision', 'conclusion', 'summary', 'action item', 'next step'
        ]
        
        sentences = text.split('.')
        for sentence in sentences:
            sentence = sentence.strip()
            if any(indicator in sentence.lower() for indicator in indicators):
                if len(sentence) > 20:  # Filter out very short sentences
                    key_points.append(sentence + '.')
                    
        return key_points[:10]  # Return top 10 key points
        
    def _extract_action_items(self, text: str) -> List[str]:
        """Extract action items from transcript"""
        action_items = []
        
        # Look for action-related phrases
        action_patterns = [
            'will do', 'need to', 'should', 'must', 'action item', 'follow up',
            'next step', 'todo', 'assign', 'responsible', 'deadline', 'by when'
        ]
        
        sentences = text.split('.')
        for sentence in sentences:
            sentence = sentence.strip()
            if any(pattern in sentence.lower() for pattern in action_patterns):
                if len(sentence) > 15:
                    action_items.append(sentence + '.')
                    
        return action_items[:8]  # Return top 8 action items
        
    def _generate_visual_summary(self, interactions: List[Dict[str, Any]]) -> str:
        """Generate summary of visual interactions"""
        if not interactions:
            return "No significant visual interactions detected."
            
        scene_changes = [i for i in interactions if i["type"] == "scene_change"]
        content_changes = [i for i in interactions if i["type"] == "content_change"]
        
        summary = f"Visual analysis detected {len(scene_changes)} scene changes and {len(content_changes)} content changes throughout the video. "
        
        if scene_changes:
            summary += f"Major scene transitions occurred at: "
            times = [i["timestamp_formatted"] for i in scene_changes[:5]]
            summary += ", ".join(times)
            if len(scene_changes) > 5:
                summary += f" and {len(scene_changes) - 5} more. "
            else:
                summary += ". "
                
        return summary
        
    def _generate_pdf_report(self, report_data: Dict[str, Any], output_dir: Path) -> str:
        """Generate PDF report using ReportLab"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.utils import ImageReader
            
            # Create output file
            output_file = output_dir / f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # Create document
            doc = SimpleDocTemplate(str(output_file), pagesize=letter, topMargin=1*inch)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.darkblue
            )
            
            # Build content
            story = []
            
            # Title
            story.append(Paragraph("Meeting Video Analysis Report", title_style))
            story.append(Spacer(1, 20))
            
            # Metadata section
            self._add_pdf_metadata_section(story, report_data, styles)
            
            # Video overview section
            self._add_pdf_video_overview_section(story, report_data, styles)
            
            # Analysis summary section
            self._add_pdf_analysis_summary_section(story, report_data, styles)
            
            # Summaries section (if available)
            if "summaries" in report_data:
                self._add_pdf_summaries_section(story, report_data, styles)
                
            # Timeline section
            self._add_pdf_timeline_section(story, report_data, styles)
            
            # Screenshots section (if enabled)
            if self.include_screenshots and report_data.get("screenshots"):
                story.append(PageBreak())
                self._add_pdf_screenshots_section(story, report_data, styles)
                
            # Full transcript section
            story.append(PageBreak())
            self._add_pdf_transcript_section(story, report_data, styles)
            
            # Build PDF
            doc.build(story)
            
            return str(output_file)
            
        except ImportError:
            self.logger.warning("ReportLab not available, falling back to text report")
            return self._generate_text_report(report_data, output_dir)
        except Exception as e:
            self.logger.error(f"PDF generation failed: {e}")
            return self._generate_text_report(report_data, output_dir)
            
    def _add_pdf_metadata_section(self, story, report_data, styles):
        """Add metadata section to PDF"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        story.append(Paragraph("Report Information", styles['Heading2']))
        
        metadata = report_data["metadata"]
        data = [
            ["Generation Time:", metadata["generation_time"]],
            ["Video File:", metadata["video_file"]],
            ["Report Format:", metadata["report_format"]],
            ["Processing Duration:", f"{metadata['processing_duration']:.2f} seconds"]
        ]
        
        table = Table(data, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Vertical alignment
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
    def _add_pdf_video_overview_section(self, story, report_data, styles):
        """Add video overview section to PDF"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        story.append(Paragraph("Video Overview", styles['Heading2']))
        
        overview = report_data["video_overview"]
        data = [
            ["Filename:", overview["filename"]],
            ["Duration:", overview["duration"]],
            ["Resolution:", overview["resolution"]],
            ["Frame Rate:", f"{overview['fps']:.2f} fps"],
            ["File Size:", overview["file_size"]],
            ["Codec:", overview["codec"]]
        ]
        
        table = Table(data, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
    def _add_pdf_analysis_summary_section(self, story, report_data, styles):
        """Add analysis summary section to PDF"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        story.append(Paragraph("Analysis Summary", styles['Heading2']))
        
        analysis = report_data["analysis_summary"]
        transcript_summary = report_data["transcript_summary"]
        
        data = [
            ["Frames Analyzed:", str(analysis["total_frames_analyzed"])],
            ["Scene Changes:", str(analysis["scene_changes_detected"])],
            ["Interactions Detected:", str(analysis["interactions_detected"])],
            ["Transcript Segments:", str(transcript_summary["total_segments"])],
            ["Total Words:", str(transcript_summary["total_words"])],
            ["Words per Minute:", f"{transcript_summary['words_per_minute']:.1f}"],
            ["Average Confidence:", f"{transcript_summary['average_confidence']:.2f}"]
        ]
        
        table = Table(data, colWidths=[2.5*inch, 3.5*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
    def _add_pdf_summaries_section(self, story, report_data, styles):
        """Add summaries section to PDF"""
        from reportlab.platypus import Paragraph, Spacer
        
        summaries = report_data.get("summaries", {})
        
        story.append(Paragraph("Content Summaries", styles['Heading2']))
        
        if "overall_summary" in summaries:
            story.append(Paragraph("Overall Summary", styles['Heading3']))
            story.append(Paragraph(summaries["overall_summary"], styles['Normal']))
            story.append(Spacer(1, 12))
            
        if "key_points" in summaries:
            story.append(Paragraph("Key Points", styles['Heading3']))
            for point in summaries["key_points"]:
                story.append(Paragraph(f"• {point}", styles['Normal']))
            story.append(Spacer(1, 12))
            
        if "action_items" in summaries:
            story.append(Paragraph("Action Items", styles['Heading3']))
            for item in summaries["action_items"]:
                story.append(Paragraph(f"• {item}", styles['Normal']))
            story.append(Spacer(1, 12))
                
    def _add_pdf_timeline_section(self, story, report_data, styles):
        """Add timeline section to PDF"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        story.append(Paragraph("Timeline of Events", styles['Heading2']))
        
        timeline = report_data["detailed_timeline"]
        
        # Create table data
        data = [[Paragraph("<b>Time</b>", styles['Normal']), 
                 Paragraph("<b>Type</b>", styles['Normal']), 
                 Paragraph("<b>Description</b>", styles['Normal'])]]
        
        for event in timeline:  # Show ALL events
            event_type = event["type"].replace("_", " ").title()
            description = event["description"]
            
            data.append([
                Paragraph(event["timestamp_formatted"], styles['Normal']),
                Paragraph(event_type, styles['Normal']),
                Paragraph(description, styles['Normal'])
            ])
            
        table = Table(data, colWidths=[1*inch, 1.5*inch, 3.5*inch])
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Vertical alignment to top
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('WORDWRAP', (0, 0), (-1, -1), True)  # Enable word wrapping
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
    def _add_pdf_screenshots_section(self, story, report_data, styles):
        """Add screenshots section to PDF"""
        from reportlab.platypus import Paragraph, Spacer, Image
        from reportlab.lib.utils import ImageReader
        from reportlab.lib.units import inch
        
        story.append(Paragraph("Key Screenshots", styles['Heading2']))
        
        screenshots = report_data.get("screenshots", [])
        
        for screenshot in screenshots[:10]:  # Limit to 10 screenshots
            if Path(screenshot["file_path"]).exists():
                try:
                    # Add timestamp and description
                    story.append(Paragraph(
                        f"Screenshot at {screenshot['timestamp_formatted']}", 
                        styles['Heading4']
                    ))
                    story.append(Paragraph(screenshot["description"], styles['Normal']))
                    
                    # Add image
                    img = Image(screenshot["file_path"], width=5*inch, height=3*inch, hAlign='CENTER')
                    story.append(img)
                    story.append(Spacer(1, 20))
                    
                except Exception as e:
                    self.logger.warning(f"Could not add screenshot {screenshot['file_path']}: {e}")
                    
    def _add_pdf_transcript_section(self, story, report_data, styles):
        """Add transcript section to PDF with properly segmented text"""
        from reportlab.platypus import Paragraph, Spacer
        import re
        
        story.append(Paragraph("Full Transcript", styles['Heading2']))
        story.append(Paragraph("(Transcript is divided into timestamped segments for easy reading)", styles['Normal']))
        story.append(Spacer(1, 12))
        
        segments = report_data.get("transcript_segments", [])
        
        for segment in segments:
            text = segment.get("text", "").strip()
            if not text:
                continue
                
            start_time = segment.get('start', 0)
            end_time = segment.get('end', 0)
            duration = end_time - start_time
            
            # Split into sentences (at . ! ?)
            sentences = re.split(r'(?<=[.!?])\s+', text)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            if len(sentences) <= 3:
                # Short segment - display as is
                timestamp_text = f"<b>[{self._format_timestamp(start_time)} - {self._format_timestamp(end_time)}]</b>"
                story.append(Paragraph(timestamp_text, styles['Normal']))
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 12))
            else:
                # Long segment - split into chunks of 2-3 sentences each
                total_sentences = len(sentences)
                sentences_per_chunk = 3
                
                for i in range(0, total_sentences, sentences_per_chunk):
                    chunk_sentences = sentences[i:i + sentences_per_chunk]
                    chunk_text = ' '.join(chunk_sentences)
                    
                    # Interpolate timestamp based on sentence position
                    chunk_start_time = start_time + (duration * i / total_sentences)
                    chunk_end_time = start_time + (duration * min(i + sentences_per_chunk, total_sentences) / total_sentences)
                    
                    timestamp_text = f"<b>[{self._format_timestamp(chunk_start_time)} - {self._format_timestamp(chunk_end_time)}]</b>"
                    story.append(Paragraph(timestamp_text, styles['Normal']))
                    story.append(Paragraph(chunk_text, styles['Normal']))
                    story.append(Spacer(1, 10))
            
    def _generate_docx_report(self, report_data: Dict[str, Any], output_dir: Path) -> str:
        """Generate Word document report"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Meeting Video Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata section
            self._add_docx_metadata_section(doc, report_data)
            
            # Video overview section
            self._add_docx_video_overview_section(doc, report_data)
            
            # Analysis summary section
            self._add_docx_analysis_summary_section(doc, report_data)
            
            # Summaries section (if available)
            if "summaries" in report_data:
                self._add_docx_summaries_section(doc, report_data)
                
            # Timeline section
            self._add_docx_timeline_section(doc, report_data)
            
            # Screenshots section (if enabled)
            if self.include_screenshots and report_data.get("screenshots"):
                self._add_docx_screenshots_section(doc, report_data)
                
            # Full transcript section
            self._add_docx_transcript_section(doc, report_data)
            
            # Save document
            output_file = output_dir / f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(str(output_file))
            
            return str(output_file)
            
        except ImportError:
            self.logger.warning("python-docx not available, falling back to text report")
            return self._generate_text_report(report_data, output_dir)
        except Exception as e:
            self.logger.error(f"DOCX generation failed: {e}")
            return self._generate_text_report(report_data, output_dir)
            
    def _add_docx_metadata_section(self, doc, report_data):
        """Add metadata section to Word document"""
        doc.add_heading('Report Information', level=1)
        
        metadata = report_data["metadata"]
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ("Generation Time:", metadata["generation_time"]),
            ("Video File:", metadata["video_file"]),
            ("Report Format:", metadata["report_format"]),
            ("Processing Duration:", f"{metadata['processing_duration']:.2f} seconds")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            
        doc.add_page_break()
        
    def _add_docx_video_overview_section(self, doc, report_data):
        """Add video overview section to Word document"""
        doc.add_heading('Video Overview', level=1)
        
        overview = report_data["video_overview"]
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ("Filename:", overview["filename"]),
            ("Duration:", overview["duration"]),
            ("Resolution:", overview["resolution"]),
            ("Frame Rate:", f"{overview['fps']:.2f} fps"),
            ("File Size:", overview["file_size"]),
            ("Codec:", overview["codec"])
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            
    def _add_docx_analysis_summary_section(self, doc, report_data):
        """Add analysis summary section to Word document"""
        doc.add_heading('Analysis Summary', level=1)
        
        analysis = report_data["analysis_summary"]
        transcript_summary = report_data["transcript_summary"]
        
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ("Frames Analyzed:", str(analysis["total_frames_analyzed"])),
            ("Scene Changes:", str(analysis["scene_changes_detected"])),
            ("Interactions Detected:", str(analysis["interactions_detected"])),
            ("Transcript Segments:", str(transcript_summary["total_segments"])),
            ("Total Words:", str(transcript_summary["total_words"])),
            ("Words per Minute:", f"{transcript_summary['words_per_minute']:.1f}"),
            ("Average Confidence:", f"{transcript_summary['average_confidence']:.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            
    def _add_docx_summaries_section(self, doc, report_data):
        """Add summaries section to Word document"""
        summaries = report_data.get("summaries", {})
        
        doc.add_heading('Content Summaries', level=1)
        
        if "overall_summary" in summaries:
            doc.add_heading('Overall Summary', level=2)
            doc.add_paragraph(summaries["overall_summary"])
            
        if "key_points" in summaries:
            doc.add_heading('Key Points', level=2)
            for point in summaries["key_points"]:
                p = doc.add_paragraph(point)
                p.style = 'List Bullet'
                
        if "action_items" in summaries:
            doc.add_heading('Action Items', level=2)
            for item in summaries["action_items"]:
                p = doc.add_paragraph(item)
                p.style = 'List Bullet'
                
    def _add_docx_timeline_section(self, doc, report_data):
        """Add timeline section to Word document"""
        doc.add_heading('Timeline of Events', level=1)
        
        timeline = report_data["detailed_timeline"]
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Time'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Description'
        
        # Add events
        for event in timeline:  # Show ALL events
            row_cells = table.add_row().cells
            row_cells[0].text = event["timestamp_formatted"]
            row_cells[1].text = event["type"].replace("_", " ").title()
            row_cells[2].text = event["description"]
            
    def _add_docx_screenshots_section(self, doc, report_data):
        """Add screenshots section to Word document"""
        doc.add_heading('Key Screenshots', level=1)
        
        screenshots = report_data.get("screenshots", [])
        
        for screenshot in screenshots[:10]:  # Limit to 10 screenshots
            if Path(screenshot["file_path"]).exists():
                try:
                    doc.add_heading(f"Screenshot at {screenshot['timestamp_formatted']}", level=3)
                    doc.add_paragraph(screenshot["description"])
                    doc.add_picture(screenshot["file_path"], width=Inches(5))
                    doc.add_paragraph()  # Add space
                except Exception as e:
                    self.logger.warning(f"Could not add screenshot {screenshot['file_path']}: {e}")
                    
    def _add_docx_transcript_section(self, doc, report_data):
        """Add transcript section to Word document with properly segmented text"""
        import re
        
        doc.add_heading('Full Transcript', level=1)
        doc.add_paragraph('(Transcript is divided into timestamped segments for easy reading)')
        
        segments = report_data.get("transcript_segments", [])
        
        for segment in segments:
            text = segment.get("text", "").strip()
            if not text:
                continue
                
            start_time = segment.get('start', 0)
            end_time = segment.get('end', 0)
            duration = end_time - start_time
            
            # Split into sentences (at . ! ?)
            sentences = re.split(r'(?<=[.!?])\s+', text)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            if len(sentences) <= 3:
                # Short segment - display as is
                timestamp_text = f"[{self._format_timestamp(start_time)} - {self._format_timestamp(end_time)}]"
                doc.add_heading(timestamp_text, level=4)
                doc.add_paragraph(text)
            else:
                # Long segment - split into chunks of 2-3 sentences each
                total_sentences = len(sentences)
                sentences_per_chunk = 3
                
                for i in range(0, total_sentences, sentences_per_chunk):
                    chunk_sentences = sentences[i:i + sentences_per_chunk]
                    chunk_text = ' '.join(chunk_sentences)
                    
                    # Interpolate timestamp based on sentence position
                    chunk_start_time = start_time + (duration * i / total_sentences)
                    chunk_end_time = start_time + (duration * min(i + sentences_per_chunk, total_sentences) / total_sentences)
                    
                    timestamp_text = f"[{self._format_timestamp(chunk_start_time)} - {self._format_timestamp(chunk_end_time)}]"
                    doc.add_heading(timestamp_text, level=4)
                    doc.add_paragraph(chunk_text)
            
    def _generate_text_report(self, report_data: Dict[str, Any], output_dir: Path) -> str:
        """Generate plain text report"""
        
        output_file = output_dir / f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Title
            f.write("=" * 60 + "\n")
            f.write("MEETING VIDEO ANALYSIS REPORT\n")
            f.write("=" * 60 + "\n\n")
            
            # Metadata
            f.write("REPORT INFORMATION\n")
            f.write("-" * 20 + "\n")
            metadata = report_data["metadata"]
            f.write(f"Generation Time: {metadata['generation_time']}\n")
            f.write(f"Video File: {metadata['video_file']}\n")
            f.write(f"Report Format: {metadata['report_format']}\n")
            f.write(f"Processing Duration: {metadata['processing_duration']:.2f} seconds\n\n")
            
            # Video Overview
            f.write("VIDEO OVERVIEW\n")
            f.write("-" * 15 + "\n")
            overview = report_data["video_overview"]
            f.write(f"Filename: {overview['filename']}\n")
            f.write(f"Duration: {overview['duration']}\n")
            f.write(f"Resolution: {overview['resolution']}\n")
            f.write(f"Frame Rate: {overview['fps']:.2f} fps\n")
            f.write(f"File Size: {overview['file_size']}\n")
            f.write(f"Codec: {overview['codec']}\n\n")
            
            # Analysis Summary
            f.write("ANALYSIS SUMMARY\n")
            f.write("-" * 17 + "\n")
            analysis = report_data["analysis_summary"]
            transcript_summary = report_data["transcript_summary"]
            f.write(f"Frames Analyzed: {analysis['total_frames_analyzed']}\n")
            f.write(f"Scene Changes: {analysis['scene_changes_detected']}\n")
            f.write(f"Interactions Detected: {analysis['interactions_detected']}\n")
            f.write(f"Transcript Segments: {transcript_summary['total_segments']}\n")
            f.write(f"Total Words: {transcript_summary['total_words']}\n")
            f.write(f"Words per Minute: {transcript_summary['words_per_minute']:.1f}\n")
            f.write(f"Average Confidence: {transcript_summary['average_confidence']:.2f}\n\n")
            
            # Summaries (if available)
            if "summaries" in report_data:
                summaries = report_data["summaries"]
                f.write("CONTENT SUMMARIES\n")
                f.write("-" * 17 + "\n")
                
                if "overall_summary" in summaries:
                    f.write("Overall Summary:\n")
                    f.write(summaries["overall_summary"] + "\n\n")
                    
                if "key_points" in summaries:
                    f.write("Key Points:\n")
                    for point in summaries["key_points"]:
                        f.write(f"• {point}\n")
                    f.write("\n")
                    
                if "action_items" in summaries:
                    f.write("Action Items:\n")
                    for item in summaries["action_items"]:
                        f.write(f"• {item}\n")
                    f.write("\n")
                    
            # Timeline
            f.write("TIMELINE OF EVENTS\n")
            f.write("-" * 18 + "\n")
            timeline = report_data["detailed_timeline"]
            
            for event in timeline:
                f.write(f"{event['timestamp_formatted']} - {event['type'].replace('_', ' ').title()}\n")
                f.write(f"  {event['description']}\n\n")
                
            # Full Transcript
            f.write("FULL TRANSCRIPT\n")
            f.write("-" * 15 + "\n")
            segments = report_data.get("transcript_segments", [])
            
            for segment in segments:
                start_time = self._format_timestamp(segment['start'])
                end_time = self._format_timestamp(segment['end'])
                f.write(f"[{start_time} - {end_time}]\n")
                f.write(f"{segment['text']}\n\n")
                
        return str(output_file)
        
    def _format_duration(self, seconds: float) -> str:
        """Format duration in human readable format"""
        if seconds < 60:
            return f"{seconds:.0f}s"
        elif seconds < 3600:
            minutes = seconds // 60
            secs = seconds % 60
            return f"{minutes:.0f}m {secs:.0f}s"
        else:
            hours = seconds // 3600
            minutes = (seconds % 3600) // 60
            secs = seconds % 60
            return f"{hours:.0f}h {minutes:.0f}m {secs:.0f}s"
            
    def _format_file_size(self, bytes_size: int) -> str:
        """Format file size in human readable format"""
        if bytes_size < 1024:
            return f"{bytes_size} B"
        elif bytes_size < 1024**2:
            return f"{bytes_size/1024:.1f} KB"
        elif bytes_size < 1024**3:
            return f"{bytes_size/(1024**2):.1f} MB"
        else:
            return f"{bytes_size/(1024**3):.1f} GB"
            
    def _format_timestamp(self, seconds: float) -> str:
        """Format timestamp as HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"